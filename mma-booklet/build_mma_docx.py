#!/usr/bin/env python3
"""
Builds "APEX PREDATOR MMA - Marketing Booklet" as an EDITABLE Word document.

Subculture: MIXED MARTIAL ARTS (MMA)

Run:  python3 build_mma_docx.py
Out:  APEX-MMA-Marketing-Booklet.docx

This is the editable companion to APEX-MMA-Marketing-Booklet.pdf. It carries the
same research copy, but as real Word headings, tables and bullet lists so any of
it can be rewritten. It uses only fonts that ship with Word, so nothing needs
installing.

Photography: every frame appears as a shaded placeholder table. Click inside one,
delete the text and use Insert > Pictures to drop your own photograph in.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "APEX-MMA-Marketing-Booklet.docx")
BRAND = os.path.join(HERE, "brand")

# ---------------------------------------------------------------- palette
RED = RGBColor(0xE6, 0x19, 0x1F)
INK = RGBColor(0x14, 0x14, 0x17)
GREY = RGBColor(0x6F, 0x70, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SILVER = RGBColor(0x8A, 0x8C, 0x92)

SH_DARK = "141417"
SH_RED = "E6191F"
SH_TINT = "EDEBE8"
SH_PH = "E4E1DD"

BODY = "Calibri"
DISP = "Impact"          # closest widely-installed match for Anton
LBLF = "Calibri"


# ---------------------------------------------------------------- xml helpers
def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement("w:" + tag)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    tblPr.append(borders)


def keep_with_next(par):
    par.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------- text helpers
def style_run(run, size=10.5, bold=False, italic=False, colour=None,
              font=BODY, caps=False, spacing=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if colour is not None:
        run.font.color.rgb = colour
    if caps:
        run.font.all_caps = True
    if spacing is not None:
        rPr = run._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(spacing * 20)))
        rPr.append(sp)
    # make sure east-asian/complex scripts use the same face
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), font)


def para(doc, text="", size=10.5, bold=False, italic=False, colour=None,
         font=BODY, before=0, after=6, align=None, line=1.15, caps=False,
         spacing=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line:
        p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        style_run(p.add_run(text), size, bold, italic, colour, font, caps, spacing)
    return p


def kicker(doc, text, before=14):
    return para(doc, text, size=8.5, bold=True, colour=RED, before=before,
                after=2, caps=True, spacing=1.2)


# Headings display as caps via Word's all-caps *formatting*, not by upper-casing
# the text itself. That keeps the underlying words normally cased, so editing a
# heading does not mean typing in caps or retyping to change the case.
def h1(doc, num, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    if num:
        style_run(p.add_run(num + "  "), 26, False, False, RED, DISP)
    style_run(p.add_run(text), 22, False, False, INK, DISP, caps=True)
    return p


def h2(doc, text, before=16):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(4)
    keep_with_next(p)
    style_run(p.add_run(text), 13, False, False, INK, DISP, caps=True)
    return p


def h3(doc, text, before=10):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(2)
    keep_with_next(p)
    style_run(p.add_run(text), 10.5, True, False, RED, BODY,
              caps=True, spacing=0.8)
    return p


def bullets(doc, items, size=10.5):
    """Bulleted list. 'Lead in - rest' renders the lead in bold."""
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        if " \u2014 " in it:
            lead, rest = it.split(" \u2014 ", 1)
            style_run(p.add_run(lead + " \u2014 "), size, True, False, INK)
            style_run(p.add_run(rest), size, False, False, INK)
        else:
            style_run(p.add_run(it), size, False, False, INK)


def rule(doc, colour=SH_RED, after=8):
    """A thin horizontal brand rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(after)
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour)
    bd.append(bottom)
    pPr.append(bd)
    return p


def run_lines(p, text, **kw):
    """Add text to a paragraph, turning \\n into real Word line breaks.

    A literal newline inside a single run is not a line break in OOXML - it has
    to be an explicit <w:br/>, or the text silently runs together.
    """
    parts = str(text).split("\n")
    for i, part in enumerate(parts):
        r = p.add_run(part)
        style_run(r, **kw)
        if i < len(parts) - 1:
            r.add_break()


def kv_table(doc, rows, widths=(3.4, 13.4)):
    """Borderless key/value rows - no header band."""
    t = doc.add_table(rows=0, cols=2)
    no_borders(t)
    t.autofit = False
    for k, v in rows:
        kc, vc = t.add_row().cells
        kc.width, vc.width = Cm(widths[0]), Cm(widths[1])
        cell_margins(kc, top=20, bottom=20, start=0, end=60)
        cell_margins(vc, top=20, bottom=20, start=0, end=0)
        style_run(kc.paragraphs[0].add_run(k), 8.5, True, False, RED, BODY,
                  caps=True, spacing=0.9)
        style_run(vc.paragraphs[0].add_run(v), 10, False, False, INK)
    return t


def table(doc, headers, rows, widths=None, head_fill=SH_DARK, first_bold=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        shade(c, head_fill)
        cell_margins(c)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        style_run(p.add_run(htxt), 8.5, True, False, WHITE, BODY,
                  caps=True, spacing=0.8)
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    for r, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            cell_margins(c)
            if r % 2 == 0:
                shade(c, SH_TINT)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.05
            bold = first_bold and i == 0
            run_lines(p, val, size=9, bold=bold,
                      colour=RED if bold else INK)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def photo_slot(doc, slot, label, brief, height_note=None):
    """A shaded placeholder the user can click into and replace with a picture."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.autofit = False
    c = t.rows[0].cells[0]
    c.width = Cm(16.8)
    shade(c, SH_PH)
    cell_margins(c, top=160, bottom=160, start=160, end=160)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run("PHOTOGRAPH \u2014 " + slot.upper() + "  \u00b7  " + label),
              9, True, False, RED, BODY, spacing=0.6)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.05
    style_run(p2.add_run("Shot brief: " + brief), 8.5, False, True, GREY)
    p3 = c.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    style_run(p3.add_run("Click this box, delete this text, then Insert > "
                         "Pictures to place your photograph here."),
              8, False, False, SILVER)
    para(doc, "", after=8)
    return t


def add_toc(doc):
    """A real Word table-of-contents field."""
    p = doc.add_paragraph()
    r = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "Right-click here and choose Update Field to build the contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (fld, instr, sep, hint, end):
        r._r.append(el)
    return p


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------- document
doc = Document()
doc.core_properties.title = "APEX PREDATOR MMA - Marketing Booklet"
doc.core_properties.author = "APEX PREDATOR MMA"
doc.core_properties.subject = "Marketing booklet - the MMA (mixed martial arts) subculture"
doc.core_properties.keywords = "MMA, subculture, semiotics, marketing, fight wear"

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.1)
sec.right_margin = Cm(2.1)

st = doc.styles["Normal"]
st.font.name = BODY
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6)

# footer
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(fp.add_run("APEX PREDATOR MMA  \u00b7  Marketing Booklet  \u00b7  "
                     "MMA subculture"), 8, False, False, GREY)

# =====================================================================
# COVER
# =====================================================================
para(doc, "MARKETING BOOKLET  \u00b7  SUBCULTURE RESEARCH  \u00b7  2026",
     size=9.5, bold=True, colour=RED, before=40, after=18, caps=True, spacing=1.6)
para(doc, "APEX PREDATOR", size=44, colour=INK, font=DISP, after=0)
para(doc, "MMA", size=64, colour=RED, font=DISP, before=0, after=6)
rule(doc, after=10)
para(doc, "FIGHT CLUB", size=22, colour=INK, font=DISP, after=0)
para(doc, "TRAINING  \u00b7  SPARRING  \u00b7  FIGHT NIGHTS", size=9.5, bold=True,
     colour=RED, after=22, caps=True, spacing=1.4)
para(doc, "A fight-wear brand built for the mixed martial arts subculture.",
     size=15, colour=INK, font=DISP, after=4)
para(doc, "Product line: FIGHT CLUB \u2014 rashguards, shorts, tees, gloves.",
     size=10.5, colour=GREY, after=26)

info = doc.add_table(rows=4, cols=2)
no_borders(info)
for i, (k, v) in enumerate((("SUBCULTURE", "MMA \u00b7 Mixed martial arts"),
                            ("BRAND", "APEX PREDATOR MMA"),
                            ("PRODUCT", "FIGHT CLUB training and fight wear"),
                            ("ISSUE", "01 / 2026"))):
    kc, vc = info.rows[i].cells
    kc.width, vc.width = Cm(4.2), Cm(12.6)
    cell_margins(kc, start=0)
    cell_margins(vc, start=0)
    style_run(kc.paragraphs[0].add_run(k), 8.5, True, False, RED, BODY,
              caps=True, spacing=1.1)
    style_run(vc.paragraphs[0].add_run(v), 10.5, False, False, INK)

para(doc, "", after=24)
para(doc, "HOW TO USE THIS DOCUMENT", size=9, bold=True, colour=RED,
     after=4, caps=True, spacing=1.1)
para(doc, "This is the editable version of the booklet. All copy can be rewritten "
          "directly. Every photograph is a shaded placeholder box showing its shot "
          "brief \u2014 click into one, delete the text and use Insert > Pictures to "
          "place your own photograph. The designed, print-ready version of the same "
          "content is APEX-MMA-Marketing-Booklet.pdf.",
     size=9.5, italic=True, colour=GREY)
page_break(doc)

# =====================================================================
# CONTENTS
# =====================================================================
h1(doc, "", "Contents")
rule(doc)
para(doc, "This booklet presents APEX PREDATOR MMA \u2014 a fight-wear brand designed "
          "and marketed to the mixed martial arts subculture. It works through the "
          "research into the group, the meaning carried by their signs and symbols, "
          "the development of the design, and the plan for selling it to them.",
     after=12)
add_toc(doc)
page_break(doc)

# =====================================================================
# 01 INTRODUCTION
# =====================================================================
kicker(doc, "Section", before=0)
h1(doc, "01", "Introduction / About")
rule(doc)
para(doc, "APEX PREDATOR MMA is a fight-wear and training-gear brand built for one "
          "group of people: the mixed martial arts subculture. Its range is called "
          "FIGHT CLUB \u2014 rashguards, shorts, tees and gloves designed to be trained "
          "in rather than posed in.")
para(doc, "The name does two jobs at once. An apex predator sits at the top of the "
          "food chain, which is exactly how a fighter is taught to think about a "
          "division. \u201cApex\u201d is also the highest point of a climb \u2014 the peak you "
          "train towards.")
para(doc, "The brand exists because fight gear sits in an awkward gap. Mainstream "
          "sportswear treats fighters as an afterthought: seams in the wrong place "
          "for grappling, sizing cut for runners, and graphics drawn by people who "
          "have never been tapped.")
para(doc, "The specialist fight labels understand the sport but often look crude \u2014 "
          "skulls, flames, barbed wire, gothic type. APEX aims at the space between "
          "the two: real technical credibility with a design language a 22-year-old "
          "would actually wear on the street.")

h2(doc, "The range")
para(doc, "Indicative RRP.", size=9, italic=True, colour=GREY, after=6)
table(doc, ("#", "Product", "RRP", "Notes"), [
    ("01", "Rashguard", "$79",
     "Sublimated print, flatlock seams, four-way stretch. The core identity piece."),
    ("02", "Fight shorts", "$69",
     "Split hem and silicone waist grip. Built to kick in without riding up."),
    ("03", "Tee / hoodie", "$45 / $95",
     "Heavyweight cotton. The off-mat, street-facing half of the brand."),
    ("04", "4oz gloves", "$89",
     "Moulded foam, thumb lock. The entry purchase for a new member."),
], widths=(1.1, 3.2, 2.0, 10.5))

para(doc, "", after=10)
photo_slot(doc, "prod-03", "PRODUCT / KIT",
           "Flat lay: gloves, wraps, mouthguard, shorts, rashguard, water bottle. "
           "Top-down, single soft source.")
photo_slot(doc, "prod-01", "PRODUCT / RASHGUARD",
           "The APEX rashguard worn and actually trained in \u2014 sweat-marked, on the "
           "mat, not on a hanger.")
page_break(doc)

# =====================================================================
# 02 DEFINITION
# =====================================================================
kicker(doc, "Subculture research", before=0)
h1(doc, "02", "Definition of the Group")
rule(doc)
para(doc, "Mixed martial arts is a full-contact combat sport that allows both "
          "striking and grappling, standing up and on the ground. It draws its "
          "techniques from boxing, Muay Thai, kickboxing, wrestling, Brazilian "
          "jiu-jitsu, judo and sambo, then tests them against each other inside a "
          "fenced cage.")
para(doc, "The MMA subculture is the community that has grown around that sport, and "
          "it is far wider than the people who actually compete. Most members will "
          "never take a professional fight. It is made up of hobbyists who turn up "
          "four nights a week, amateur competitors, coaches and referees, and fans "
          "who will happily argue about a guard pass for an hour. What separates it "
          "from ordinary sports fandom is that membership is proven by participation "
          "\u2014 you are in it because you train, not because you bought a ticket. It "
          "should not be confused with gym or bodybuilding culture, which is "
          "organised around how a body looks rather than what it can do, nor with "
          "traditional martial arts, where forms and grading can matter more than "
          "live sparring.")

h2(doc, "What makes it different from the wider culture")
bullets(doc, [
    "Consented violence \u2014 the wider culture is organised around avoiding physical "
    "confrontation. This one deliberately seeks it, under strict rules, and treats "
    "it as the most honest test of a person available.",
    "Earned status \u2014 no amount of money, followers or talk moves you up. A blue "
    "belt beats a beginner every time, and everyone in the room knows it.",
    "Public losing \u2014 a hobbyist is submitted by training partners several times a "
    "week. Elsewhere failure is hidden; here it is timetabled, witnessed and "
    "treated as information.",
    "Closeness with strangers \u2014 grappling needs a level of body contact the wider "
    "culture reserves for intimacy or medicine. On the mat it is completely "
    "unremarkable.",
    "Chosen hardship \u2014 weight cuts, 6am runs, ice baths, no alcohol during camp. "
    "Discomfort is not endured here, it is sought out on purpose.",
    "The body keeps the record \u2014 cauliflower ear, taped fingers, scar tissue. "
    "Members carry permanent, visible proof of belonging.",
])
para(doc, "", after=8)
photo_slot(doc, "def-01", "GRAPPLING",
           "Two athletes in a clinch or a guard pass. Get low, fill the frame, keep "
           "both faces partly visible.")
photo_slot(doc, "def-02", "STRIKING",
           "Pad work with sweat spray caught in the light. Fast shutter (1/500s) and "
           "a high ISO.")
page_break(doc)

# =====================================================================
# 03 DEMOGRAPHICS
# =====================================================================
kicker(doc, "Subculture research", before=0)
h1(doc, "03", "Demographics of the Group")
rule(doc)
para(doc, "Who is actually in the room. The figures below are indicative estimates "
          "drawn from gym-level observation, promotion audience reporting and "
          "participation data rather than a formal survey, so they are given as "
          "ranges.")

table(doc, ("Indicative profile", "Figure"), [
    ("Core training age", "18\u201329 (broader active range 16\u201334)"),
    ("Male participation", "~75% of gym membership and audience"),
    ("Training frequency", "4\u20136 sessions per week"),
    ("Global reach", "40+ countries running professional cards"),
], widths=(6.0, 10.8))
para(doc, "", after=10)

for head, q, txt in [
    ("Age", "What age group are they mostly?",
     "The training population is young. The core sits at 18\u201329, inside a broader "
     "active range of 16\u201334. Under-16s arrive through kids' and teens' programs, "
     "now a major part of most gyms' income, and masters divisions keep competitors "
     "going well past 35. Spectatorship skews a little older than participation, "
     "roughly 18\u201344, because watching costs far less time than training does. "
     "APEX targets 16\u201328: old enough to buy their own gear, young enough to still "
     "be building an identity around the gym."),
    ("Gender", "Is it mainly male, female, or mixed?",
     "Still majority male \u2014 around three-quarters of gym membership and a similar "
     "share of the audience. But women's participation is the fastest-growing part "
     "of the sport, pushed along by athletes such as Ronda Rousey, Amanda Nunes, "
     "Zhang Weili and Alexa Grasso, and by women entering through self-defence and "
     "fitness classes rather than competition. Most fight labels still treat women's "
     "kit as a shrunken, pinker version of the men's. That is an obvious gap, and "
     "APEX is built to fill it."),
    ("Nationality", "Where are they from?",
     "Global, with clear strongholds: the United States, Brazil, Russia and "
     "Dagestan, Japan, Thailand, Ireland, England, Poland, Mexico, Kazakhstan, and "
     "Australia and New Zealand. In Australia the sport is concentrated in the outer "
     "suburbs of Sydney, Melbourne, Brisbane and Perth, and Australian gyms punch "
     "well above their weight internationally. APEX is positioned as an Australian "
     "brand that still reads as globally credible."),
    ("Cultural identity", "What cultural backgrounds or influences do they have?",
     "MMA is multicultural by construction: its techniques come from Brazil, Japan, "
     "Thailand, the United States and Russia, and the gym's vocabulary borrows with "
     "them \u2014 oss, mongkhon, kimura, sambo. Australian gyms are among the most "
     "mixed rooms you will find, with Pacific Islander, M\u0101ori, Lebanese, "
     "Vietnamese, Sudanese and Anglo-Australian athletes on one mat. Dagestani and "
     "wider Muslim influence has made prayer and fasting normal in elite camps, and "
     "the culture keeps deep working-class roots \u2014 for many families the gym is a "
     "route to discipline, belonging and occasionally a career."),
]:
    h2(doc, head, before=12)
    para(doc, q, size=9.5, italic=True, colour=GREY, after=4)
    para(doc, txt)

para(doc, "", after=8)
photo_slot(doc, "demo-01", "THE CLASS",
           "The whole class lined up at the end of session \u2014 deliberately show the "
           "mixed ages, genders and backgrounds in the room.")
photo_slot(doc, "demo-02", "WOMEN TRAIN",
           "Female athlete drilling or sparring, framed exactly like the male "
           "athletes: as an athlete, never as decoration.")
page_break(doc)

# =====================================================================
# 04 INTERESTS, ACTIVITIES & VALUES
# =====================================================================
kicker(doc, "Subculture research", before=0)
h1(doc, "04", "Interests, Activities & Values")
rule(doc)
para(doc, "What the members of this subculture actually have in common \u2014 the things "
          "they do together, and the ideas they agree on. These are the hooks a brand "
          "can legitimately hang itself on.")

h2(doc, "Hobbies & activities")
bullets(doc, [
    "Training \u2014 four to six sessions a week across striking, grappling and "
    "strength, plus open mats at the weekend.",
    "Competing \u2014 amateur \u201csmokers\u201d, BJJ tournaments, Muay Thai interclubs and "
    "local fight cards.",
    "Watching together \u2014 UFC and ONE cards are social events, watched in groups at "
    "odd hours because of time zones.",
    "Analysis \u2014 breaking fights down frame by frame, technique channels, podcasts "
    "and coaches' breakdowns.",
    "Gear \u2014 comparing gloves, shin pads, rashguards and mouthguards is a hobby in "
    "its own right.",
    "Recovery \u2014 ice baths, saunas, physio, macro tracking and managing weight.",
])

h2(doc, "Values & beliefs")
bullets(doc, [
    "Respect \u2014 touch gloves, shake hands, thank your partner, and protect them in "
    "sparring.",
    "Leave your ego at the door \u2014 being beaten is data, not humiliation.",
    "Discipline over motivation \u2014 turning up when you do not feel like it is the "
    "entire skill.",
    "Meritocracy \u2014 rank is demonstrated, never claimed. The mats settle it.",
    "Controlled aggression \u2014 violence has a place, a time and a rule set. Fighters "
    "are often the calmest people outside it.",
    "Loyalty and lineage \u2014 you represent a gym and a coach, and you say whose "
    "student you are.",
    "Self-mastery \u2014 the real opponent is your own comfort.",
])

para(doc, "", after=10)
qt = doc.add_table(rows=1, cols=1)
qt.style = "Table Grid"
qc = qt.rows[0].cells[0]
qc.width = Cm(16.8)
shade(qc, SH_DARK)
cell_margins(qc, top=200, bottom=200, start=200, end=200)
qp = qc.paragraphs[0]
qp.paragraph_format.space_after = Pt(2)
style_run(qp.add_run("\u201cLeave your ego at the door.\u201d"), 20, False, False,
          WHITE, DISP)
qp2 = qc.add_paragraph()
style_run(qp2.add_run("The most repeated line in any MMA gym."), 9, True, False,
          RED, BODY, caps=True, spacing=1.1)
para(doc, "", after=10)

photo_slot(doc, "belief-02", "RESPECT",
           "The glove touch or the bow at the start of a round \u2014 the gesture that "
           "turns violence into sport.")
page_break(doc)

# =====================================================================
# 05 LIFESTYLE & ATTITUDES
# =====================================================================
kicker(doc, "Subculture research", before=0)
h1(doc, "05", "Lifestyle & Attitudes")
rule(doc)
para(doc, "How belonging to this subculture actually shapes a week, a body and a set "
          "of opinions. This is the part a brand has to get right, because the "
          "audience can smell an outsider instantly.")
bullets(doc, [
    "Built around the timetable \u2014 meals, sleep and social life bend to the "
    "training schedule, not the other way around.",
    "Sacrifice is normal \u2014 diet, alcohol and late nights go first, and weight cuts "
    "before competition are brutal and accepted.",
    "Pain is reframed \u2014 soreness is expected, injuries become stories, and "
    "strapping tape is everyday equipment.",
    "Function over fashion on the mat \u2014 but strong loyalty to labels seen as "
    "legitimate, and instant contempt for the rest.",
    "Zero tolerance for fakes \u2014 poseurs, gym-only tough guys and McDojo black "
    "belts are the culture's favourite target. Authenticity is the highest currency "
    "there is.",
    "The gym as therapy \u2014 one of the most commonly stated reasons for training is "
    "managing stress, anger and anxiety.",
])

h2(doc, "A training week \u2014 the committed hobbyist")
table(doc, ("Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"), [
    ("Run 6am\nMuay Thai", "S&C\nBJJ (gi)", "MMA sparring", "Run 6am\nMuay Thai",
     "BJJ (no-gi)", "Open mat", "Recovery"),
], widths=(2.4,) * 7, first_bold=False)
para(doc, "Six training slots across seven days, with one deliberate rest day. The "
          "whole week is built around this grid.",
     size=9, italic=True, colour=GREY, before=4)

para(doc, "", after=8)
photo_slot(doc, "belief-01", "THE COACH",
           "Coach correcting a student's stance mid-round, hand on shoulder. "
           "Documentary, unposed.")
page_break(doc)

# =====================================================================
# 06 SEMIOTICS: SIGNS & SYMBOLS
# =====================================================================
kicker(doc, "Semiotic investigation", before=0)
h1(doc, "06", "Signs & Symbols")
rule(doc)
para(doc, "The signifiers this subculture uses to recognise its own members, and what "
          "each one actually means to somebody inside the gym. Almost none of these "
          "meanings are decorative \u2014 they are claims about how much work a person "
          "has done.")

h3(doc, "Denotation")
para(doc, "The literal meaning \u2014 what the thing plainly is, before any "
          "interpretation.", after=6)
h3(doc, "Connotation", before=4)
para(doc, "The implied meaning \u2014 what it represents to the group that uses it.",
     after=10)

table(doc, ("Signifier", "Denotation \u2014 what it is",
            "Connotation \u2014 what it means to them"), [
    ("The octagon / the cage",
     "An eight-sided competition area enclosed by chain-link fencing.",
     "The ultimate proving ground. Nowhere to hide and no way out, so it reads as a "
     "complete test of a person \u2014 close to sacred ground."),
    ("4oz open-finger gloves",
     "Minimal padding across the knuckles with the fingers left free.",
     "The real thing. Free fingers mean grappling is allowed, so this is a whole "
     "fight rather than just boxing."),
    ("Wrapped hands",
     "Cotton or elastic wrap supporting the wrist and knuckles.",
     "Ritual and preparation \u2014 the deliberate few minutes before violence. Reads "
     "as craft, not rage, which is why it is the most useful image the culture owns."),
    ("The BJJ belt",
     "A coloured cloth belt, running white through to black.",
     "Verified time served. A black belt is a decade of work, and it is the one "
     "credential here that cannot be bought, inherited or talked into existence."),
    ("Cauliflower ear",
     "Scar tissue in the outer ear caused by repeated friction and impact.",
     "Irrefutable proof of real mat mileage. Worn as a badge of honour, and often "
     "left untreated on purpose."),
    ("Team rashguard & gym patch",
     "A compression top printed with a gym name, crest or sponsor marks.",
     "Affiliation, exactly like a football jersey. It announces who you represent "
     "and whose student you are."),
    ("Predator imagery",
     "Skulls, wolves, big cats, snakes and claw marks.",
     "Killer instinct and dominance \u2014 sitting at the top of the food chain in your "
     "division. This is the signifier APEX PREDATOR MMA is built on."),
    ("Black + blood red",
     "A near-black base carrying a single saturated red accent.",
     "Black for discipline, seriousness and fight night. Red for blood, danger and "
     "adrenaline. Brushed silver for hardness and steel."),
    ("The glove touch / bow",
     "A brief physical gesture exchanged before and after a round.",
     "Mutual consent and respect. It is the signal that converts an assault into a "
     "sport, and both athletes know it."),
    ("Kanji, Thai & Portuguese",
     "Non-English words and characters on gear and gym walls.",
     "Respect paid to where the techniques came from. Read as lineage and "
     "authenticity rather than decoration \u2014 but only if spelled correctly."),
], widths=(3.4, 5.6, 7.8))

para(doc, "", after=8)
photo_slot(doc, "semio-01", "THE BELT",
           "Close crop on a worn BJJ belt being tied, frayed and faded. The wear is "
           "the whole point.")
photo_slot(doc, "semio-02", "THE CAGE",
           "Cage mesh close-up with a fighter soft-focused behind it. Shoot through "
           "the fence, do not shoot the fence.")
page_break(doc)

# =====================================================================
# 07 SEMIOTICS: MEANING & MISREADING
# =====================================================================
kicker(doc, "Semiotic investigation", before=0)
h1(doc, "07", "Meaning & Misreading")
rule(doc)
para(doc, "Every signifier on the previous page carries a second, unintended reading "
          "for people outside the subculture. This matters commercially as well as "
          "ethically: the same image that earns respect on the mat can lose a "
          "stockist, a sponsor or a parent's permission outside it.")

table(doc, ("Signifier", "Read inside the subculture",
            "Misread by everyone else"), [
    ("The cage",
     "A safety structure that stops fighters falling out of the ring, and the "
     "sport's proving ground.",
     "\u201cHuman cockfighting\u201d \u2014 the phrase a US senator used against the sport in "
     "the 1990s. Animals in a pen; barbarism."),
    ("Cauliflower ear",
     "A credential. Evidence of years of honest work on the mat.",
     "Disfigurement, or proof that somebody is a thug. Often the first thing a "
     "non-training parent notices."),
    ("Predator & skull imagery",
     "Competitive mindset and dominance within a weight division.",
     "Glorifying real violence \u2014 and in some contexts read as gang or extremist "
     "coding."),
    ("The name \u201cFight Club\u201d",
     "The ordinary, everyday word for a training group at a gym.",
     "The Fincher film: illegal, unsupervised, bare-knuckle fighting."),
    ("The name \u201cApex Predator\u201d",
     "Top of the food chain, competitively \u2014 the athlete nobody can solve.",
     "\u201cPredatory\u201d, which is a genuinely unhelpful word to have circling a brand "
     "aimed partly at teenagers."),
    ("Black + red, tattoos, shaved heads",
     "The standard, unremarkable aesthetic of fight culture.",
     "Intimidating; occasionally confused with football-hooligan or far-right "
     "visual coding."),
    ("Women training",
     "Athletes doing exactly what the men in the room are doing.",
     "Easily read as sexualised if it is shot using the conventions of fashion "
     "rather than those of sport."),
], widths=(3.4, 6.2, 7.2))

h2(doc, "What this means for the design")
bullets(doc, [
    "Keep the octagon, lose the blood \u2014 the cage is used as geometry, a crest and "
    "a frame, never as a picture of somebody being hurt.",
    "No skulls \u2014 the predator idea is carried by the claw slash and the word APEX "
    "alone.",
    "Photograph the discipline, not the damage \u2014 wraps, coaching, the empty gym at "
    "6am. Effort reads as aspirational to insiders and outsiders alike.",
    "Frame FIGHT CLUB as sport every time \u2014 it never appears without TRAINING \u00b7 "
    "SPARRING \u00b7 FIGHT NIGHTS beneath it, so the underground-fighting reading never "
    "lands.",
    "Shoot women exactly like men \u2014 same lenses, same angles, same sweat, same "
    "crops.",
    "Get the technique right \u2014 a wrong guard or a bad stance in an advert is "
    "mocked instantly, and the credibility does not come back.",
])

para(doc, "", after=8)
tt = doc.add_table(rows=1, cols=1)
tt.style = "Table Grid"
tc = tt.rows[0].cells[0]
tc.width = Cm(16.8)
shade(tc, SH_DARK)
cell_margins(tc, top=200, bottom=200, start=200, end=200)
tp = tc.paragraphs[0]
tp.paragraph_format.space_after = Pt(4)
style_run(tp.add_run("THE TEST APPLIED TO EVERY SYMBOL IN THIS BRAND"), 9, True,
          False, RED, BODY, caps=True, spacing=1.2)
tp2 = tc.add_paragraph()
style_run(tp2.add_run("If it only works inside the gym it is a liability. "
                      "If it only works outside, it is invisible."),
          15, False, False, WHITE, DISP)
page_break(doc)

# =====================================================================
# 08 MOOD BOARD
# =====================================================================
kicker(doc, "The product / design", before=0)
h1(doc, "08", "Mood Board")
rule(doc)
para(doc, "The visual territory the brand is aiming at: low light, hard rim light, "
          "worn equipment, steel and chalk, and effort rather than injury. Assembled "
          "before any layout work began.")

poster = os.path.join(BRAND, "APEX-MMA-poster.png")
if os.path.exists(poster):
    h2(doc, "The result \u2014 cover artwork")
    doc.add_picture(poster, height=Cm(11.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, "The finished APEX PREDATOR MMA cover artwork.", size=9,
         italic=True, colour=GREY)

h2(doc, "Mood palette")
table(doc, ("Swatch", "Hex", "Reads as"), [
    ("Night", "#08080A", "The base. Fight night, the dark outside the lit ring."),
    ("Steel", "#141417", "Panels and raised surfaces. Hardness."),
    ("Old blood", "#7C0A0D", "Shadow tone. History, previous fights."),
    ("Fight night", "#E6191F", "The single accent. Danger and adrenaline."),
    ("Adrenaline", "#FF5A52", "Highlight only. Heat and rim light."),
    ("Chalk", "#C9CCD2", "Type and detail. Brushed steel, chalk dust."),
], widths=(3.6, 3.0, 10.2))
para(doc, "The palette was pulled out of the photographs rather than chosen first.",
     size=9, italic=True, colour=GREY, before=4)

para(doc, "", after=8)
for s, l, br in (("mood-01", "MOOD / SWEAT",
                  "High-contrast black and white of a fighter's back and shoulders, "
                  "sweat catching the light."),
                 ("mood-02", "MOOD / NIGHT",
                  "Arena or gym lights flaring in the dark \u2014 atmosphere, no "
                  "subject."),
                 ("mood-03", "MOOD / STILL",
                  "A fighter sitting alone on the mat between rounds, head down, "
                  "breathing. The quiet side of the sport.")):
    photo_slot(doc, s, l, br)
page_break(doc)

# =====================================================================
# 09 IDEATION
# =====================================================================
kicker(doc, "The product / design", before=0)
h1(doc, "09", "Ideation & Development")
rule(doc)
para(doc, "Turning the research into something that can be drawn. The mind map "
          "converts each researched signifier into a usable design device; the "
          "thumbnails test how those devices sit on a page.")

h2(doc, "Mind map \u2014 from research to design device")
table(doc, ("Researched signifier", "Becomes this design device"), [
    ("The octagon", "Crest shape, page frame, and an eight-sided grid."),
    ("Wraps & tape", "The hero photograph, and a repeating texture."),
    ("Sweat & grain", "Duotone treatment and film grain over photography."),
    ("Claw / predator", "The three-slash mark inside the logo."),
    ("Black + blood red", "The palette: near-black with exactly one accent."),
    ("Lineage & respect", "Small-caps labelling and visible credits."),
], widths=(5.4, 11.4))

h2(doc, "Thumbnail sketches")
table(doc, ("#", "Layout", "Idea", "Outcome"), [
    ("01", "Cage frame", "Octagon border with the type set inside it.", "Rejected \u2014 too enclosed."),
    ("02", "Split type", "Red / silver two-tone headline over a photo.", "SELECTED"),
    ("03", "Full bleed", "Photo edge to edge with a type band across it.", "Rejected \u2014 type got lost."),
    ("04", "Crest lockup", "Centred crest, fully symmetrical.", "Rejected \u2014 too static."),
    ("05", "Slab stack", "Three stacked type slabs.", "Rejected \u2014 no room for photography."),
], widths=(1.1, 3.2, 7.5, 5.0))

h2(doc, "Brainstorm")
para(doc, "APEX  \u00b7  cage  \u00b7  OCTAGON  \u00b7  wraps  \u00b7  CLAW  \u00b7  lineage  \u00b7  SWEAT  "
          "\u00b7  oss  \u00b7  PREDATOR  \u00b7  blood red  \u00b7  tap  \u00b7  DISCIPLINE  \u00b7  mat time  "
          "\u00b7  HONOR  \u00b7  steel  \u00b7  cauliflower  \u00b7  WARRIOR  \u00b7  6am",
     size=12, colour=INK)
page_break(doc)

# =====================================================================
# 10 DESIGN SYSTEM
# =====================================================================
kicker(doc, "The product / design", before=0)
h1(doc, "10", "The Design System")
rule(doc)
para(doc, "The finished kit of parts, and the research reason behind each decision. "
          "Every element here traces back to a signifier identified in sections 06 "
          "and 07.")

h2(doc, "Logo / lockup")
bullets(doc, [
    "The crest \u2014 a true eight-sided octagon, taken straight from the cage. It is "
    "the one shape the subculture cannot mistake for anything else.",
    "The claw \u2014 three slashes, the centre one longest. Carries the predator idea "
    "without resorting to a skull.",
    "The split \u2014 APEX PREDATOR in silver, MMA in red, so the sport is always the "
    "loudest word in the lockup.",
])

h2(doc, "Colour palette")
table(doc, ("Name", "Hex", "Used for"), [
    ("Ink", "#08080A", "Base / background"),
    ("Steel", "#141417", "Panels"),
    ("Deep red", "#7C0A0D", "Shadow"),
    ("Accent red", "#E6191F", "The single accent"),
    ("Highlight red", "#FF5A52", "Highlights only"),
    ("Silver", "#C9CCD2", "Type on dark"),
], widths=(4.2, 3.4, 9.2))

h2(doc, "Typefaces")
table(doc, ("Face", "Role", "Why"), [
    ("Anton", "Display, headlines, the lockup",
     "Heavy, condensed and unapologetic. Behaves like stencilled arena signage and "
     "holds up at poster scale."),
    ("Barlow Condensed", "Labels, small caps, captions",
     "Narrow enough to set the long technical strings the sport is full of without "
     "shouting over the display face."),
], widths=(4.0, 4.6, 8.2))
para(doc, "In this Word version, Impact and Calibri stand in for Anton and Barlow "
          "Condensed so the document opens correctly on any machine. Install Anton "
          "and Barlow Condensed (both free, from Google Fonts) and set the headings "
          "to them if you want an exact match with the PDF.",
     size=9, italic=True, colour=GREY, before=4)

h2(doc, "Applications")
table(doc, ("Application", "How the system is applied"), [
    ("Rashguard", "Crest centred on the chest, one red band at the hem."),
    ("Tee / hoodie", "Crest only, small, left chest. Street-readable."),
    ("Poster", "Split headline, one red accent bar, full-bleed photograph."),
    ("Instagram", "Square crops, crest watermark, alternating red and near-black tiles."),
], widths=(4.2, 12.6))
para(doc, "Every application is built from the same three parts: the crest, the split "
          "wordmark, and exactly one red accent.",
     size=9, italic=True, colour=GREY, before=4)
page_break(doc)

# =====================================================================
# 11 PHOTOGRAPHY PLAN
# =====================================================================
kicker(doc, "Production", before=0)
h1(doc, "11", "Photography Plan")
rule(doc)
para(doc, "Photography is a required element of this product, and it is also the "
          "single biggest credibility test the brand faces. Every frame is shot on "
          "location in a working gym, using people who actually train.")

h2(doc, "How these are shot")
bullets(doc, [
    "Available light first \u2014 one window or a single LED panel, with the background "
    "deliberately underexposed so the subject is rim-lit rather than evenly lit.",
    "Fast glass, high ISO \u2014 1/500s and above to freeze strikes; 1/30s deliberately "
    "for the wide room shots so movement smears and the gym reads as busy.",
    "Documentary, not directed \u2014 nothing is staged twice. Sweat, tape and scuffed "
    "gear stay in frame because removing them removes the proof.",
    "Consent and safety \u2014 written permission from the gym and from every athlete, "
    "and nobody is photographed mid-injury.",
])

h2(doc, "Shot list \u2014 all 20 frames")
SHOTS = [
    ("hero-01", "Hero / cover", "Fighter mid-roundhouse on the pads, shot from low and slightly behind, hard rim light, gym dark behind. Underexpose the background by 2 stops."),
    ("strip-01", "Hand wraps", "Tight crop on hands being wrapped \u2014 the ritual before training. Side light, shallow depth of field."),
    ("strip-02", "Glove detail", "Worn 4oz gloves resting on the cage or the mat edge. Show the scuffs; do not clean them up."),
    ("strip-03", "Mat texture", "Overhead of the tatami seams and chalk-dusted mat, feet just entering frame."),
    ("gym-01", "The room", "Wide of the whole gym mid-class, bodies blurred by a slow shutter (1/30s) so the room reads busy and alive."),
    ("def-01", "Grappling", "Two athletes in a clinch or a guard pass. Get low, fill the frame, keep both faces partly visible."),
    ("def-02", "Striking", "Pad work with sweat spray caught in the light. Fast shutter (1/500s) and a high ISO."),
    ("demo-01", "The class", "The whole class lined up at the end of session \u2014 deliberately show the mixed ages, genders and backgrounds."),
    ("demo-02", "Women train", "Female athlete drilling or sparring, framed exactly like the male athletes: as an athlete, never as decoration."),
    ("belief-01", "The coach", "Coach correcting a student's stance mid-round, hand on shoulder. Documentary, unposed."),
    ("belief-02", "Respect", "The glove touch or the bow at the start of a round \u2014 the gesture that turns violence into sport."),
    ("semio-01", "The belt", "Close crop on a worn BJJ belt being tied, frayed and faded. The wear is the whole point."),
    ("semio-02", "The cage", "Cage mesh close-up with a fighter soft-focused behind it. Shoot through the fence, do not shoot the fence."),
    ("mood-01", "Mood / sweat", "High-contrast black and white of a fighter's back and shoulders, sweat catching the light."),
    ("mood-02", "Mood / night", "Arena or gym lights flaring in the dark \u2014 atmosphere, no subject."),
    ("mood-03", "Mood / still", "A fighter sitting alone on the mat between rounds, head down, breathing."),
    ("prod-01", "Product / rashguard", "The APEX rashguard worn and actually trained in \u2014 sweat-marked, on the mat, not on a hanger."),
    ("prod-02", "Product / tee", "The FIGHT CLUB tee worn outside the gym, street context, so it reads as everyday clothing too."),
    ("prod-03", "Product / kit", "Flat lay: gloves, wraps, mouthguard, shorts, rashguard, water bottle. Top-down, single soft source."),
    ("market-01", "The walkout", "Fighter walking out under the lights with the brand visible on the kit."),
]
table(doc, ("#", "Frame", "Save as", "Brief"),
      [(str(i), lbl, code + ".jpg", br)
       for i, (code, lbl, br) in enumerate(SHOTS, 1)],
      widths=(0.9, 3.4, 2.9, 9.6))
page_break(doc)

# =====================================================================
# 12 MARKETABILITY
# =====================================================================
kicker(doc, "Marketability", before=0)
h1(doc, "12", "Who Would Buy This")
rule(doc)
para(doc, "Being specific about the target audience matters more here than in almost "
          "any other market, because this subculture actively polices who is allowed "
          "to claim membership. The three buyers below are ranked by how much "
          "credibility they carry, not by how much they spend.")

for nm, age, role, flag, facts, note in [
    ("Tyson", "19", "The committed hobbyist", "PRIMARY TARGET",
     [("Trains", "5 nights a week, 18 months in"),
      ("Level", "One amateur Muay Thai bout"),
      ("Buys", "2\u20133 rashguards a year, shorts, mouthguards"),
      ("Spend", "$250\u2013400 a year")],
     "Wants gear that says he is serious rather than brand new. He is the buyer "
     "whose approval everyone else copies."),
    ("Aisha", "23", "The competitor", "HIGHEST VALUE",
     [("Trains", "6 sessions a week plus S&C"),
      ("Level", "BJJ blue belt, competes quarterly"),
      ("Buys", "Competition-legal kit, walkout wear"),
      ("Spend", "$400\u2013700 a year")],
     "Judges fit and performance before anything else. A brand that finally cuts "
     "properly for women earns loyalty that is very hard to shift."),
    ("Jack", "17", "The fan at the edge", "HIGHEST VOLUME",
     [("Trains", "Once or twice a week"),
      ("Level", "No competition, watches every card"),
      ("Buys", "Tees and hoodies only"),
      ("Spend", "$80\u2013150 a year")],
     "Buys the look and the belonging. Essential for volume, but must never be "
     "served at the expense of the two buyers above."),
]:
    h2(doc, "%s, %s \u2014 %s" % (nm, age, role))
    para(doc, flag, size=8.5, bold=True, colour=RED, after=4, caps=True,
         spacing=1.2)
    kv_table(doc, facts)
    para(doc, note, size=10, italic=True, colour=GREY, before=4)

h2(doc, "Two more things the plan depends on")
bullets(doc, [
    "Gyms are the real customer \u2014 a single team-kit order for one gym outsells "
    "dozens of individual sales, and it puts the logo on a whole room at once. This "
    "is the biggest revenue lever the brand has.",
    "Who we are deliberately not chasing \u2014 general activewear buyers, and anyone "
    "drawn to street-fighting imagery. Serving either one would cost the credibility "
    "that everything else here depends on.",
])

para(doc, "", after=8)
photo_slot(doc, "market-01", "THE WALKOUT",
           "Fighter walking out under the lights with the brand visible on the kit "
           "\u2014 the moment the logo gets broadcast.")
page_break(doc)

# =====================================================================
# 13 MARKETING PLAN
# =====================================================================
kicker(doc, "Marketability", before=0)
h1(doc, "13", "How We Reach Them")
rule(doc)
para(doc, "This audience does not respond to conventional advertising, and it is "
          "unusually good at spotting a brand run by people who do not train. The "
          "plan therefore spends most of its effort inside gyms and at events, and "
          "treats social media as documentation rather than promotion.")

h2(doc, "Channels")
table(doc, ("#", "Channel", "What we actually do"), [
    ("01", "Instagram & TikTok",
     "Short technique clips, sparring rounds, gear checks and fight-camp footage. "
     "Athlete-shot and vertical, never studio-polished. Posted daily."),
    ("02", "Athlete & gym seeding",
     "Free kit to local amateurs and coaches in exchange for walkout and training "
     "wear. Cheap, and it buys the one thing money normally cannot: insider "
     "endorsement."),
    ("03", "Events & interclubs",
     "A merch booth and cage banner at local fight nights, BJJ comps and Muay Thai "
     "interclubs. Sponsor one card per quarter."),
    ("04", "Gym stockists",
     "Wholesale racks inside partner gyms plus bulk team-kit deals, so the brand is "
     "bought at the exact place identity is formed."),
    ("05", "Online store",
     "Shopify, with an honest fit guide by discipline, and limited drops to create "
     "the scarcity this audience already responds to."),
    ("06", "Long-form & community",
     "YouTube breakdowns, podcast sponsorship, and a genuine presence in the sport's "
     "forums and Discords rather than paid comments."),
], widths=(1.1, 4.4, 11.3))

h2(doc, "Budget split \u2014 year one")
table(doc, ("Line", "Share", "Reasoning"), [
    ("Athlete & gym seeding", "28%", "Buys credibility that cannot be bought any other way."),
    ("Photography & content", "24%", "The brand is carried almost entirely by its images."),
    ("Events & sponsorship", "20%", "Physical presence where the audience already gathers."),
    ("Paid social", "16%", "Amplification only, never the primary channel."),
    ("Store & sampling", "12%", "Fit guides, samples, returns and size exchanges."),
], widths=(5.4, 2.4, 9.0))

h2(doc, "The three rules")
bullets(doc, [
    "Use real athletes from real gyms \u2014 a model who has never trained is spotted "
    "in one frame.",
    "Never fake credentials \u2014 no invented lineage, no borrowed belts, no stolen "
    "fight records.",
    "Support before selling \u2014 sponsor the local card, pay the photographer, kit "
    "out the coach.",
])
page_break(doc)

# =====================================================================
# 14 CONCLUSION
# =====================================================================
kicker(doc, "Section", before=0)
h1(doc, "14", "Conclusion")
rule(doc)
para(doc, "APEX PREDATOR MMA takes a subculture that is built on proof rather than "
          "talk, and answers it with a brand that can survive being looked at "
          "closely.")

for hd, tx in (("What the research found",
                "A community organised around earned status, chosen hardship and a "
                "total intolerance of anything fake."),
               ("What the symbols mean",
                "The octagon, the wraps, the belt and the worn-in gear all say the "
                "same thing: this person has done the work."),
               ("What the design does",
                "Uses those signs as geometry and restraint rather than as gore, so "
                "it reads as credible inside the gym and safe outside it.")):
    h3(doc, hd)
    para(doc, tx)

para(doc, "", after=14)
para(doc, "WARRIOR", size=30, colour=RED, font=DISP, after=0)
para(doc, "SPIRIT", size=30, colour=INK, font=DISP, before=0, after=12)

ct = doc.add_table(rows=1, cols=1)
ct.style = "Table Grid"
cc = ct.rows[0].cells[0]
cc.width = Cm(16.8)
shade(cc, SH_RED)
cell_margins(cc, top=180, bottom=180, start=200, end=200)
cp = cc.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_after = Pt(2)
style_run(cp.add_run("JOIN THE ELITE   |   APEX FIGHT CLUB"), 18, False, False,
          WHITE, DISP)
cp2 = cc.add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(cp2.add_run("TRAINING \u00b7 SPARRING \u00b7 FIGHT NIGHTS"), 9, True, False,
          WHITE, BODY, caps=True, spacing=2.0)

para(doc, "", after=10)
para(doc, "WWW.APEXMMA.COM", size=15, colour=INK, font=DISP,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, "@APEXMMA   \u00b7   #UNLEASHTHEBEAST", size=9, bold=True, colour=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, caps=True, spacing=1.8, after=18)

rule(doc)
para(doc, "Note on sources: the demographic figures in section 03 are indicative "
          "estimates drawn from gym-level observation and published audience "
          "profiles, not survey data. Prices are indicative RRP for a fictional "
          "brand.",
     size=8.5, italic=True, colour=GREY)

doc.save(OUT)
print("Wrote %s" % os.path.basename(OUT))
print("  sections: 14   photo placeholders: %d   shot list rows: %d"
      % (14, len(SHOTS)))
